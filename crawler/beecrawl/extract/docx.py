"""DOCX text extraction (IHBB sample rounds ship as .docx)."""

from __future__ import annotations

import io

import docx as python_docx


def extract_docx(data: bytes) -> str:
    document = python_docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs]

    # Some sample rounds lay questions out in a two-column table
    # (question | answer), which the paragraph walk misses entirely.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(p.rstrip() for p in parts).strip()
